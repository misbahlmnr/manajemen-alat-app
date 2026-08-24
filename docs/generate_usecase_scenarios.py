# -*- coding: utf-8 -*-
"""Generator dokumen Word untuk Use Case Scenario Sistem Manajemen Alat & Bahan Lab.

Menghasilkan 10 tabel skenario use case sesuai use-case-diagram.drawio.
Jalankan: python generate_usecase_scenarios.py
"""

from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


USE_CASES = [
    {
        "no": 1,
        "name": "Login",
        "id": "UC01",
        "importance": "High",
        "actor": "Admin, Guru, Siswa",
        "type": "Main Use Case",
        "stakeholder": "Pengguna ingin masuk ke sistem sesuai hak akses perannya masing-masing.",
        "brief": "Use Case ini menjelaskan bagaimana pengguna melakukan proses login sebelum dapat menggunakan fitur sesuai perannya.",
        "trigger": "Pengguna ingin membuka halaman beranda sistem.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "-",
        "extend": "Logout (UC02)",
        "generalization": "-",
        "normal": [
            "Pengguna membuka halaman login sistem.",
            "Pengguna memasukkan email dan kata sandi, lalu menekan tombol masuk.",
            "Sistem memeriksa data yang dimasukkan pengguna.",
            "Sistem membuka halaman beranda sesuai peran (Admin, Guru, atau Siswa).",
        ],
        "exceptional": "Jika email atau kata sandi salah, sistem menampilkan pesan gagal masuk dan pengguna tetap berada di halaman login.",
    },
    {
        "no": 2,
        "name": "Logout",
        "id": "UC02",
        "importance": "Medium",
        "actor": "Admin, Guru, Siswa",
        "type": "Main Use Case",
        "stakeholder": "Pengguna ingin keluar dari sistem dengan aman setelah selesai menggunakannya.",
        "brief": "Use Case ini menjelaskan bagaimana pengguna mengakhiri sesi dan keluar dari sistem.",
        "trigger": "Pengguna ingin mengakhiri penggunaan sistem.",
        "uc_type": "Primary",
        "relationship": "Extend (dari Login / UC01)",
        "include": "-",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Pengguna menekan menu keluar pada sistem.",
            "Sistem mengakhiri sesi pengguna.",
            "Sistem mengarahkan pengguna kembali ke halaman login.",
        ],
        "exceptional": "Jika sesi pengguna sudah berakhir, sistem langsung menampilkan halaman login.",
    },
    {
        "no": 3,
        "name": "Mengelola Master Data",
        "id": "UC03",
        "importance": "High",
        "actor": "Admin",
        "type": "Main Use Case",
        "stakeholder": "Admin ingin menjaga data alat, bahan, pengguna, jadwal praktikum, dan jaminan kartu tetap akurat.",
        "brief": "Use Case ini menjelaskan bagaimana admin menambah, mengubah, dan menghapus data utama pada sistem (alat, bahan, pengguna, jadwal praktikum, dan jaminan kartu).",
        "trigger": "Admin perlu memperbarui data utama sistem.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Admin membuka menu data yang ingin dikelola (Alat, Bahan, Pengguna, Jadwal, atau Jaminan Kartu).",
            "Sistem menampilkan daftar data yang ada.",
            "Admin menambah, mengubah, atau menghapus data.",
            "Admin menyimpan perubahan.",
            "Sistem menyimpan dan menampilkan data terbaru.",
        ],
        "exceptional": "Jika isian tidak lengkap atau tidak sesuai, sistem menampilkan pesan kesalahan dan data tidak disimpan.",
    },
    {
        "no": 4,
        "name": "Verifikasi Peminjaman & Permintaan Bahan",
        "id": "UC04",
        "importance": "High",
        "actor": "Admin",
        "type": "Main Use Case",
        "stakeholder": "Admin ingin memverifikasi pengajuan peminjaman alat dan permintaan bahan dari siswa.",
        "brief": "Use Case ini menjelaskan bagaimana admin menyetujui atau menolak pengajuan peminjaman alat maupun permintaan bahan, serta mengatur prioritas antrian bila stok terbatas.",
        "trigger": "Ada pengajuan peminjaman alat atau permintaan bahan yang masuk.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Admin membuka menu Peminjaman.",
            "Sistem menampilkan daftar pengajuan peminjaman alat dan permintaan bahan yang menunggu persetujuan.",
            "Admin memeriksa detail pengajuan.",
            "Admin menyetujui atau menolak pengajuan.",
            "Sistem memperbarui status dan memberi tahu siswa.",
        ],
        "exceptional": "Jika pengajuan ditolak, admin mengisi alasan penolakan. Jika stok masih kurang, pengajuan berada di antrian dan admin dapat mengatur prioritas antrian.",
    },
    {
        "no": 5,
        "name": "Pengembalian Alat",
        "id": "UC05",
        "importance": "High",
        "actor": "Admin, Siswa",
        "type": "Main Use Case",
        "stakeholder": "Siswa ingin mengembalikan alat, dan admin ingin memastikan alat kembali dalam kondisi baik.",
        "brief": "Use Case ini menjelaskan proses pengembalian alat, mulai dari pengajuan oleh siswa hingga inspeksi kondisi dan pencatatan oleh admin, termasuk penanganan jaminan kartu bila diperlukan.",
        "trigger": "Alat yang dipinjam akan dikembalikan.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Siswa mengajukan pengembalian alat melalui menu peminjaman & permintaan.",
            "Admin menerima pengajuan pengembalian dari siswa.",
            "Admin memeriksa kondisi alat yang dikembalikan (inspeksi).",
            "Admin mencatat pengembalian, lalu sistem menambah kembali stok alat.",
            "Sistem mengubah status peminjaman menjadi Dikembalikan.",
        ],
        "exceptional": "Jika alat rusak atau hilang, admin mencatat hasil inspeksi dan menandainya perlu kompensasi. Untuk alat yang dibawa pulang, kartu jaminan dikembalikan setelah alat diterima.",
    },
    {
        "no": 6,
        "name": "Membuat Laporan",
        "id": "UC06",
        "importance": "Medium",
        "actor": "Admin, Guru",
        "type": "Main Use Case",
        "stakeholder": "Admin dan guru ingin memperoleh rekap peminjaman, permintaan bahan, dan inventaris untuk keperluan pelaporan.",
        "brief": "Use Case ini menjelaskan bagaimana admin atau guru menyusun dan mengunduh laporan peminjaman alat, permintaan bahan, serta inventaris.",
        "trigger": "Admin atau guru membutuhkan laporan.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Pengguna membuka menu Laporan.",
            "Pengguna memilih jenis laporan dan rentang tanggal yang diinginkan.",
            "Sistem menampilkan data laporan sesuai pilihan.",
            "Pengguna mengunduh atau mencetak laporan.",
        ],
        "exceptional": "Jika tidak ada data pada rentang yang dipilih, sistem menampilkan laporan kosong.",
    },
    {
        "no": 7,
        "name": "Monitoring Peminjaman & Permintaan Siswa",
        "id": "UC07",
        "importance": "Medium",
        "actor": "Guru",
        "type": "Main Use Case",
        "stakeholder": "Guru ingin memantau peminjaman alat dan permintaan bahan oleh siswa, termasuk riwayatnya.",
        "brief": "Use Case ini menjelaskan bagaimana guru memantau aktivitas peminjaman alat dan permintaan bahan siswa, melihat detail pengajuan, serta meninjau riwayat yang terkait.",
        "trigger": "Guru ingin memantau kegiatan peminjaman atau permintaan bahan siswa.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Guru membuka menu Peminjaman Siswa.",
            "Sistem menampilkan daftar peminjaman alat dan permintaan bahan siswa yang berkaitan dengan guru tersebut.",
            "Guru dapat membuka tampilan riwayat untuk melihat pengajuan yang sudah selesai.",
            "Guru memilih salah satu data untuk melihat detail.",
            "Sistem menampilkan detail pengajuan beserta jadwal praktikum terkait.",
        ],
        "exceptional": "Jika belum ada data peminjaman atau permintaan bahan, sistem menampilkan pesan bahwa data belum tersedia.",
    },
    {
        "no": 8,
        "name": "Lihat Inventaris",
        "id": "UC08",
        "importance": "Medium",
        "actor": "Guru, Siswa",
        "type": "Main Use Case",
        "stakeholder": "Guru dan siswa ingin mengetahui ketersediaan alat dan stok bahan laboratorium.",
        "brief": "Use Case ini menjelaskan bagaimana guru atau siswa melihat katalog inventaris alat dan bahan lab beserta status ketersediaannya.",
        "trigger": "Pengguna ingin mengecek ketersediaan alat atau bahan lab.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Pengguna membuka menu Alat Lab atau Bahan Lab.",
            "Sistem menampilkan daftar alat atau bahan beserta jumlah tersedia dan statusnya.",
            "Pengguna memilih salah satu barang untuk melihat detailnya.",
            "Sistem menampilkan detail barang, termasuk kondisi atau informasi stok.",
        ],
        "exceptional": "Jika barang yang dicari tidak ditemukan, sistem menampilkan pesan bahwa data tidak tersedia.",
    },
    {
        "no": 9,
        "name": "Mengajukan Peminjaman & Permintaan Bahan",
        "id": "UC09",
        "importance": "High",
        "actor": "Siswa",
        "type": "Main Use Case",
        "stakeholder": "Siswa ingin meminjam alat atau meminta bahan untuk keperluan praktikum maupun penggunaan di lab.",
        "brief": "Use Case ini menjelaskan bagaimana siswa mengajukan peminjaman alat atau permintaan bahan kepada admin. Bahan diajukan sebagai permintaan (bukan peminjaman).",
        "trigger": "Siswa ingin meminjam alat atau meminta bahan.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Siswa membuka menu Ajukan Alat / Bahan.",
            "Siswa memilih jenis pengajuan, yaitu peminjaman alat atau permintaan bahan.",
            "Siswa memilih barang dan jumlahnya, serta mengisi data pendukung (misalnya keperluan, ruang pemakaian, atau jadwal praktikum).",
            "Siswa menekan tombol kirim.",
            "Sistem menyimpan pengajuan dan menampilkan status Menunggu Persetujuan.",
        ],
        "exceptional": "Jika stok tidak mencukupi, sistem memasukkan pengajuan ke antrian dan memberi tahu status antrian kepada siswa. Jika data belum lengkap, sistem menampilkan pesan agar siswa melengkapi isian.",
    },
    {
        "no": 10,
        "name": "Kelola Peminjaman & Permintaan Saya",
        "id": "UC10",
        "importance": "Medium",
        "actor": "Siswa",
        "type": "Main Use Case",
        "stakeholder": "Siswa ingin memantau dan mengelola pengajuan peminjaman alat maupun permintaan bahannya, termasuk melihat riwayat.",
        "brief": "Use Case ini menjelaskan bagaimana siswa melihat daftar aktif dan riwayat, mengubah atau membatalkan pengajuan yang masih menunggu, serta mengajukan pengembalian alat yang sedang dipinjam.",
        "trigger": "Siswa ingin melihat, mengubah, atau meninjau status peminjaman alat maupun permintaan bahannya.",
        "uc_type": "Primary",
        "relationship": "Association",
        "include": "Login (UC01)",
        "extend": "-",
        "generalization": "-",
        "normal": [
            "Siswa membuka menu Alat & Bahan Saya.",
            "Sistem menampilkan daftar peminjaman alat dan permintaan bahan yang aktif.",
            "Siswa dapat membuka tampilan riwayat untuk melihat pengajuan yang sudah selesai.",
            "Siswa memilih salah satu data untuk melihat detail.",
            "Siswa dapat mengubah atau membatalkan pengajuan yang masih menunggu, atau mengajukan pengembalian alat yang sedang dipinjam.",
            "Sistem memperbarui status pengajuan.",
        ],
        "exceptional": "Jika pengajuan sudah disetujui atau diproses, pilihan ubah dan batal tidak lagi tersedia. Permintaan bahan tidak melalui proses pengembalian seperti peminjaman alat.",
    },
]


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "left", "bottom", "right"):
        element = OxmlElement(f"w:{edge}")
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "4")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), "000000")
        tcBorders.append(element)


def clear_cell(cell):
    cell.text = ""
    return cell.paragraphs[0]


def add_label_value(paragraph, label, value, label_italic=True):
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.space_before = Pt(2)
    run_label = paragraph.add_run(label)
    run_label.italic = label_italic
    run_label.bold = False
    if value:
        run_value = paragraph.add_run(value)
        run_value.italic = False
    return paragraph


def add_stacked(cell, label, value):
    p = clear_cell(cell)
    add_label_value(p, label, "")
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    r = p2.add_run(value)
    r.bold = False
    set_cell_border(cell)


def build_table(document, uc):
    document.add_paragraph()
    caption = document.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = caption.add_run(f"Tabel 3.{uc['no']} Use Case Scenario {uc['name']}")
    cap_run.italic = True
    cap_run.bold = True

    table = document.add_table(rows=0, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = True

    row = table.add_row().cells
    add_stacked(row[0], "Use Case Name:", uc["name"])
    add_stacked(row[1], "ID:", uc["id"])
    p = clear_cell(row[2])
    add_label_value(p, "Importance: ", uc["importance"])
    set_cell_border(row[2])

    row = table.add_row().cells
    merged = row[0].merge(row[1])
    p = clear_cell(merged)
    add_label_value(p, "Primary Actor: ", uc["actor"])
    set_cell_border(merged)
    p = clear_cell(row[2])
    add_label_value(p, "Use Case Type: ", uc["type"])
    set_cell_border(row[2])

    row = table.add_row().cells
    merged = row[0].merge(row[1]).merge(row[2])
    p = clear_cell(merged)
    add_label_value(p, "Stakeholder and Interest:", "")
    p2 = merged.add_paragraph()
    p2.add_run(uc["stakeholder"])
    set_cell_border(merged)

    row = table.add_row().cells
    merged = row[0].merge(row[1]).merge(row[2])
    p = clear_cell(merged)
    add_label_value(p, "Brief Description:", "")
    p2 = merged.add_paragraph()
    p2.add_run(uc["brief"])
    set_cell_border(merged)

    row = table.add_row().cells
    merged = row[0].merge(row[1]).merge(row[2])
    p = clear_cell(merged)
    add_label_value(p, "Trigger: ", uc["trigger"])
    for label, value in [
        ("Type: ", uc["uc_type"]),
        ("Relationship: ", uc["relationship"]),
        ("Include: ", uc["include"]),
        ("Extend: ", uc["extend"]),
        ("Generalization/Inheritance: ", uc["generalization"]),
    ]:
        pp = merged.add_paragraph()
        add_label_value(pp, label, value)
    set_cell_border(merged)

    row = table.add_row().cells
    merged = row[0].merge(row[1]).merge(row[2])
    p = clear_cell(merged)
    add_label_value(p, "Normal Flow of Events:", "")
    for i, step in enumerate(uc["normal"], start=1):
        pp = merged.add_paragraph()
        pp.paragraph_format.left_indent = Pt(18)
        pp.add_run(f"{i}. {step}")
    set_cell_border(merged)

    row = table.add_row().cells
    merged = row[0].merge(row[1]).merge(row[2])
    p = clear_cell(merged)
    add_label_value(p, "Exceptional Flows:", "")
    p2 = merged.add_paragraph()
    p2.add_run(uc["exceptional"])
    set_cell_border(merged)

    return table


def main():
    document = Document()

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    title = document.add_heading("Use Case Scenario", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Sistem Manajemen Alat & Bahan Laboratorium")
    sub_run.bold = True

    intro = document.add_paragraph(
        "Berikut adalah skenario dari 10 use case pada Sistem Manajemen Alat & Bahan "
        "Laboratorium yang menjelaskan interaksi antara aktor (Admin, Guru, dan Siswa) "
        "dengan sistem sesuai use case diagram."
    )
    intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for uc in USE_CASES:
        build_table(document, uc)

    output_path = Path(__file__).resolve().parent / "usecase-scenario.docx"
    document.save(output_path)
    print(f"Dokumen berhasil dibuat: {output_path}")


if __name__ == "__main__":
    main()
